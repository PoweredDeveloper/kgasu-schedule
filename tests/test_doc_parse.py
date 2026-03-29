"""Tests for schedule text extraction and parsing."""

from __future__ import annotations

import io

from docx import Document

from scraper.doc_parse import (
    WEEKDAYS,
    _hhmm_from_bitrix_time,
    docx_extract,
    parse_cells,
    parse_schedule_from_antiword,
    plain_text_from_schedule_file,
)


def test_hhmm_from_bitrix_time_single_digit_hour() -> None:
    assert _hhmm_from_bitrix_time("9:40:00") == "9:40"
    assert _hhmm_from_bitrix_time("09:40:00") == "9:40"


def test_parse_cells_strips_edges() -> None:
    assert parse_cells("| a | b |") == ["a", "b"]
    assert parse_cells("a|b|c|d") == ["a", "b", "c", "d"]


def test_parse_schedule_docx_style_row_with_repeated_weekday() -> None:
    text = (
        "Понедельник | 1 | Чет | 9:40:00 - 11:10:00 | Математика\n"
        "Понедельник | 2 | Чет | 11:45:00 - 13:15:00 | Физика\n"
    )
    sched = parse_schedule_from_antiword(text)
    assert len(sched["Monday"]) == 2
    assert sched["Monday"][0]["time"] == "9:40-11:10"
    assert "Математика" in sched["Monday"][0]["subject"]


def test_parse_schedule_english_weekday_column() -> None:
    text = "Monday | 1 | Wk | 10:00:00 - 11:00:00 | Test subject\n"
    sched = parse_schedule_from_antiword(text)
    assert len(sched["Monday"]) == 1
    assert sched["Monday"][0]["subject"] == "Test subject"


def test_docx_extract_paragraph_and_table() -> None:
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Header line")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    doc.save(buf)
    raw = buf.getvalue()
    out = docx_extract(raw)
    assert "Header line" in out
    assert "A | B" in out


def test_plain_text_from_schedule_file_prefers_docx_for_pk_bytes() -> None:
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("From docx")
    doc.save(buf)
    raw = buf.getvalue()
    text = plain_text_from_schedule_file("https://x/file.docx", raw)
    assert "From docx" in text


def test_plain_text_from_schedule_file_empty_bytes() -> None:
    assert plain_text_from_schedule_file("https://x/y.docx", b"") == ""


def test_parse_schedule_empty_text() -> None:
    sched = parse_schedule_from_antiword("")
    assert sched == {d: [] for d in WEEKDAYS}
