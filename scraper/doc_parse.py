"""Extract structured lessons from antiword plain text (KGASU .doc tables)."""

from __future__ import annotations

import logging
import os
import re
import subprocess
import tempfile
from typing import Any

logger = logging.getLogger(__name__)

WEEKDAYS = [
    "Monday",
    "Tuesday",
    "Wednesday",
    "Thursday",
    "Friday",
    "Saturday",
    "Sunday",
]

WEEKDAYS_EN_SET = frozenset(WEEKDAYS)

RU_DAY_TO_EN: dict[str, str] = {
    "понедельник": "Monday",
    "вторник": "Tuesday",
    "среда": "Wednesday",
    "четверг": "Thursday",
    "пятница": "Friday",
    "суббота": "Saturday",
    "воскресенье": "Sunday",
}


def _norm_day(s: str) -> str:
    return s.replace("ё", "е").replace("Ё", "е").lower()


def _hhmm_from_bitrix_time(t: str) -> str:
    """Normalize '9:40:00' to '9:40' ([:5] breaks single-digit hours)."""
    parts = t.strip().split(":")
    if len(parts) >= 2:
        return f"{int(parts[0])}:{parts[1].zfill(2)}"
    return t.strip()[:5]


def _empty_week() -> dict[str, list[dict[str, str]]]:
    return {d: [] for d in WEEKDAYS}


def parse_cells(line: str) -> list[str]:
    raw = line.rstrip("\n\r")
    if "|" not in raw:
        return []
    parts = [p.strip() for p in raw.split("|")]
    if parts and parts[0] == "":
        parts = parts[1:]
    if parts and parts[-1] == "":
        parts = parts[:-1]
    return parts


def parse_schedule_from_antiword(text: str) -> dict[str, list[dict[str, str]]]:
    """Parse antiword output into weekday -> lessons."""
    by_day = _empty_week()
    current_day: str | None = None
    day_acc = ""
    pending: dict[str, Any] | None = None

    def flush() -> None:
        nonlocal pending
        if not pending or current_day is None:
            pending = None
            return
        subj = " ".join(pending.get("subj", [])).strip()
        start = pending.get("start")
        end = pending.get("end")
        if not subj or not start:
            pending = None
            return
        if end:
            time_s = f"{start}-{end}"
        else:
            time_s = start
        by_day[current_day].append(
            {
                "time": time_s,
                "subject": subj,
                "teacher": "",
                "room": "",
            }
        )
        pending = None

    time_full = re.compile(
        r"^(\d{1,2}:\d{2}:\d{2})\s*-\s*(\d{1,2}:\d{2}:\d{2})"
    )
    time_start_only = re.compile(r"^(\d{1,2}:\d{2}:\d{2})\s*-\s*$")
    time_only = re.compile(r"^(\d{1,2}:\d{2}:\d{2})$")

    for line in text.splitlines():
        if "|" not in line:
            continue
        cells = parse_cells(line)
        if len(cells) < 4:
            continue
        c0, c3 = cells[0], cells[3]
        tail = " ".join(cells[4:]).strip()

        t3 = c3.strip()
        m_full = time_full.match(t3)
        m_start = time_start_only.match(t3)
        is_lesson_time = m_full is not None or m_start is not None

        # .docx tables repeat the weekday in column 0 on every row; only treat as
        # «day header» when column 3 is not a time slot (legacy antiword layout).
        if is_lesson_time and c0:
            wdt = c0.strip().title()
            if wdt in WEEKDAYS_EN_SET:
                flush()
                current_day = wdt
                day_acc = ""
            elif re.search(r"[а-яА-ЯёЁ]", c0):
                nd = _norm_day(c0.replace(" ", ""))
                if nd in RU_DAY_TO_EN:
                    flush()
                    current_day = RU_DAY_TO_EN[nd]
                    day_acc = ""
                else:
                    for ru, en in RU_DAY_TO_EN.items():
                        if ru.startswith(nd):
                            flush()
                            current_day = en
                            day_acc = ""
                            break
        elif c0 and re.search(r"[а-яА-ЯёЁ]", c0):
            flush()
            day_acc += c0.replace(" ", "")
            nd = _norm_day(day_acc)
            if nd in RU_DAY_TO_EN:
                current_day = RU_DAY_TO_EN[nd]
                day_acc = ""
            elif any(ru.startswith(nd) for ru in RU_DAY_TO_EN):
                pass
            else:
                day_acc = c0.replace(" ", "")
            continue
        elif c0 and c0.strip().title() in WEEKDAYS_EN_SET:
            flush()
            current_day = c0.strip().title()
            day_acc = ""
            continue

        if current_day is None:
            continue

        if pending and tail and not re.search(r"\d{1,2}:\d{2}:\d{2}", c3):
            pending.setdefault("subj", []).append(tail)
            continue

        if m_full:
            flush()
            s, e = _hhmm_from_bitrix_time(m_full.group(1)), _hhmm_from_bitrix_time(m_full.group(2))
            pending = {"start": s, "end": e, "subj": []}
            if tail:
                pending["subj"].append(tail)
                flush()
            continue

        if m_start:
            flush()
            pending = {"start": _hhmm_from_bitrix_time(m_start.group(1)), "end": None, "subj": []}
            if tail:
                pending["subj"].append(tail)
            continue

        if pending and time_only.match(t3) and not tail:
            pending["end"] = _hhmm_from_bitrix_time(time_only.match(t3).group(1))
            flush()
            continue

    flush()
    return by_day


def docx_extract(doc_bytes: bytes) -> str:
    """Pull plain text from Word .docx (paragraphs + tables) for the same parser as antiword."""
    import io

    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx is required for .docx schedules (pip install python-docx)")
        return ""
    if not doc_bytes or doc_bytes[:2] != b"PK":
        return ""
    try:
        doc = Document(io.BytesIO(doc_bytes))
        lines: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                lines.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells: list[str] = []
                seen_cell: set[str] = set()
                for c in row.cells:
                    ct = (c.text or "").strip()
                    if ct and ct not in seen_cell:
                        seen_cell.add(ct)
                        cells.append(ct)
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as e:
        logger.warning("docx_extract failed: %s", e)
        return ""


def plain_text_from_schedule_file(url: str, data: bytes) -> str:
    """Route .docx (ZIP) vs legacy .doc (OLE) to the right extractor."""
    if not data:
        return ""
    u = url.lower().split("?", 1)[0]
    if u.endswith(".docx") or data[:2] == b"PK":
        text = docx_extract(data)
        if text.strip():
            return text
    return antiword_extract(data)


def antiword_extract(doc_bytes: bytes) -> str:
    """Run antiword on .doc bytes. Returns empty string if antiword is missing or fails."""
    if not doc_bytes:
        return ""
    antiword_bin = os.environ.get("ANTIWORD_PATH", "antiword")
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(doc_bytes)
            path = tmp.name
        try:
            proc = subprocess.run(
                [antiword_bin, path],
                capture_output=True,
                text=True,
                timeout=120,
                check=False,
            )
            if proc.returncode != 0:
                logger.warning("antiword exit %s: %s", proc.returncode, proc.stderr[:200])
            return proc.stdout or ""
        finally:
            try:
                os.unlink(path)
            except OSError:
                pass
    except FileNotFoundError:
        logger.error("antiword not found — install it for structured schedules (Docker image includes it)")
        return ""
